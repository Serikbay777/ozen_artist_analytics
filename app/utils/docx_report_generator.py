"""
Генератор DOCX отчетов для артистов
Простая замена плейсхолдеров в шаблоне
"""

import logging
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Словарь для конвертации месяцев на русский
MONTHS_RU = {
    'January': 'января', 'February': 'февраля', 'March': 'марта',
    'April': 'апреля', 'May': 'мая', 'June': 'июня',
    'July': 'июля', 'August': 'августа', 'September': 'сентября',
    'October': 'октября', 'November': 'ноября', 'December': 'декабря'
}


class DocxReportGenerator:
    """Генератор отчетов на основе DOCX шаблонов"""
    
    def __init__(self, analytics_service):
        self.analytics = analytics_service
        self.base_dir = Path(__file__).parent.parent.parent
        self.templates_dir = self.base_dir / "reports" / "templates"
        self.reports_dir = self.base_dir / "reports" / "artist_reports"
        
        # Создаем папки если их нет
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.reports_dir.mkdir(parents=True, exist_ok=True)
    
    @staticmethod
    def _set_table_borders(table):
        """Добавляет рамки к таблице"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Создаем borders элемент
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def generate_report(
        self,
        artist_name: str,
        period: str,
        template_name: str = None,
        eur_to_kzt_rate: float = 520.0,
        copyright_percentage: float = 0.08,
    ):
        """
        Генерирует отчет на основе DOCX шаблона
        
        Args:
            artist_name: Имя артиста
            period: Период отчета
            template_name: Имя файла шаблона
            eur_to_kzt_rate: Курс EUR → KZT
            copyright_percentage: Процент авторских отчислений
        """
        try:
            logger.info(f"📊 Генерация отчета для {artist_name}")
            
            # 1. Собираем данные артиста
            df = self.analytics._df.copy()
            artist_mask = df['Исполнитель'].str.lower().str.contains(artist_name.lower(), na=False)
            artist_df = df[artist_mask].copy()
            
            if len(artist_df) == 0:
                return {
                    "success": False,
                    "error": f"Артист '{artist_name}' не найден"
                }
            
            # Точное имя артиста
            exact_name = artist_df['Исполнитель'].mode()[0]
            
            # Конвертируем числа
            if artist_df['Количество'].dtype == 'object':
                artist_df['Количество'] = artist_df['Количество'].astype(str).str.replace(',', '.').astype(float)
            if artist_df['Сумма вознаграждения'].dtype == 'object':
                artist_df['Сумма вознаграждения'] = artist_df['Сумма вознаграждения'].astype(str).str.replace(',', '.').astype(float)
            
            # Считаем суммы
            total_revenue_eur = artist_df['Сумма вознаграждения'].sum()
            total_streams = int(artist_df['Количество'].sum())
            
            # Конвертируем в тенге
            distribution_payment = int(total_revenue_eur * eur_to_kzt_rate)
            copyright_payment = int(distribution_payment * copyright_percentage)
            total_payment = distribution_payment + copyright_payment
            
            # Топ треки для релизов
            top_tracks = (
                artist_df.groupby('Название трека')
                .agg({'Количество': 'sum', 'Сумма вознаграждения': 'sum'})
                .sort_values('Количество', ascending=False)
                .head(5)
            )
            releases_list = ", ".join(top_tracks.index.tolist())
            
            # 2. Загружаем шаблон
            if template_name is None:
                template_name = "ozen_template_final.docx"
            
            template_path = self.templates_dir / template_name
            
            if not template_path.exists():
                return {
                    "success": False,
                    "error": f"Шаблон {template_name} не найден"
                }
            
            doc = Document(str(template_path))
            logger.info(f"✓ Загружен шаблон: {template_name}")
            
            # 3. Заменяем плейсхолдеры в шаблоне
            # Форматируем дату на русском
            date_eng = datetime.now().strftime('%d %B %Y года')
            month_eng = datetime.now().strftime('%B')
            date_ru = date_eng.replace(month_eng, MONTHS_RU.get(month_eng, month_eng))
            
            replacements = {
                '17 ноября 2025 года': date_ru,
                '{{ARTIST_NAME}}': exact_name,
                '{{RELEASES}}': releases_list,
                '{{PERIOD}}': period,
                '{{DISTRIBUTION_PAYMENT}': f"{distribution_payment:,}".replace(',', ' ') + ' тенге',  # Добавляем тенге так как в шаблоне нет
                '{{DISTRIBUTION_PAYMENT}}': f"{distribution_payment:,}".replace(',', ' ') + ' тенге',
                '{{COPYRIGHT_PAYMENT}}': f"{copyright_payment:,}".replace(',', ' ') + ' тенге',
                '{{TOTAL_PAYMENT}}': f"{total_payment:,}".replace(',', ' ') + ' тенге',
            }
            
            # Заменяем текст в параграфах
            for para in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in para.text:
                        # Если текст найден в параграфе, заменяем весь параграф
                        # Сохраняем форматирование первого run
                        full_text = para.text
                        new_full_text = full_text.replace(old_text, str(new_text))
                        
                        if new_full_text != full_text:
                            # Очищаем все runs
                            for run in para.runs:
                                run.text = ''
                            # Добавляем новый текст в первый run
                            if para.runs:
                                para.runs[0].text = new_full_text
                            else:
                                para.add_run(new_full_text)
            
            # Заменяем в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if old_text in run.text:
                                            run.text = run.text.replace(old_text, str(new_text))
            
            # 4. Добавляем вторую страницу с аналитикой
            doc.add_page_break()
            
            # Заголовок страницы
            p = doc.add_paragraph()
            run = p.add_run(f'Детальная Аналитика — {exact_name}')
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True
            
            p = doc.add_paragraph()
            run = p.add_run(f'Период: {period}')
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            doc.add_paragraph()
            
            # 1. ТОП-5 ТРЕКОВ
            p = doc.add_paragraph()
            run = p.add_run('Топ-5 Треков')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            
            top_5_tracks = top_tracks.sort_values('Количество', ascending=False).head(5)
            table1 = doc.add_table(rows=len(top_5_tracks) + 1, cols=3)
            self._set_table_borders(table1)
            
            # Заголовки
            hdr = table1.rows[0].cells
            hdr[0].text = 'Название трека'
            hdr[1].text = 'Стримы'
            hdr[2].text = 'Доход (EUR)'
            
            # Данные
            for i, (track_name, row) in enumerate(top_5_tracks.iterrows()):
                cells = table1.rows[i + 1].cells
                cells[0].text = track_name
                cells[1].text = f"{int(row['Количество']):,}".replace(',', ' ')
                cells[2].text = f"€{row['Сумма вознаграждения']:,.2f}"
            
            doc.add_paragraph()
            
            # 2. ТОП-10 ПЛАТФОРМ
            p = doc.add_paragraph()
            run = p.add_run('Топ-10 Платформ по Доходу')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            
            top_platforms = (
                artist_df.groupby('Платформа')
                .agg({'Сумма вознаграждения': 'sum', 'Количество': 'sum'})
                .sort_values('Сумма вознаграждения', ascending=False)
                .head(10)
            )
            
            table2 = doc.add_table(rows=len(top_platforms) + 1, cols=4)
            self._set_table_borders(table2)
            
            # Заголовки
            hdr = table2.rows[0].cells
            hdr[0].text = 'Платформа'
            hdr[1].text = 'Доход (EUR)'
            hdr[2].text = 'Стримы'
            hdr[3].text = '% от общего'
            
            # Данные
            for i, (platform, row) in enumerate(top_platforms.iterrows()):
                cells = table2.rows[i + 1].cells
                cells[0].text = platform
                cells[1].text = f"€{row['Сумма вознаграждения']:,.2f}"
                cells[2].text = f"{int(row['Количество']):,}".replace(',', ' ')
                percentage = (row['Сумма вознаграждения'] / total_revenue_eur) * 100
                cells[3].text = f"{percentage:.1f}%"
            
            doc.add_paragraph()
            
            # 3. ТОП-5 СТРАН
            p = doc.add_paragraph()
            run = p.add_run('Топ-5 Стран по Доходу')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            
            top_countries = (
                artist_df.groupby('страна / регион')
                .agg({'Сумма вознаграждения': 'sum', 'Количество': 'sum'})
                .sort_values('Сумма вознаграждения', ascending=False)
                .head(5)
            )
            
            table3 = doc.add_table(rows=len(top_countries) + 1, cols=4)
            self._set_table_borders(table3)
            
            # Заголовки
            hdr = table3.rows[0].cells
            hdr[0].text = 'Страна'
            hdr[1].text = 'Доход (EUR)'
            hdr[2].text = 'Стримы'
            hdr[3].text = '% от общего'
            
            # Данные
            for i, (country, row) in enumerate(top_countries.iterrows()):
                cells = table3.rows[i + 1].cells
                cells[0].text = country
                cells[1].text = f"€{row['Сумма вознаграждения']:,.2f}"
                cells[2].text = f"{int(row['Количество']):,}".replace(',', ' ')
                percentage = (row['Сумма вознаграждения'] / total_revenue_eur) * 100
                cells[3].text = f"{percentage:.1f}%"
            
            # 5. Сохраняем
            safe_name = exact_name.replace(' ', '_').replace('/', '_')
            output_name = f"{safe_name}_Report_{period.replace(' ', '_')}.docx"
            output_path = self.reports_dir / output_name
            
            doc.save(str(output_path))
            logger.info(f"✅ Отчет сохранен: {output_path}")
            
            # 6. Формируем ответ
            summary = (
                f"📊 Отчет для {exact_name}\n"
                f"📅 Период: {period}\n"
                f"💰 Доход: €{total_revenue_eur:,.2f} / {total_payment:,} ₸\n"
                f"🎵 Стримы: {total_streams:,}\n"
                f"🎤 Треков: {len(top_tracks)}"
            )
            
            return {
                "success": True,
                "artist_name": exact_name,
                "file_path": str(output_path),
                "file_name": output_name,
                "summary": summary
            }
            
        except Exception as e:
            logger.error(f"❌ Ошибка: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e)
            }

